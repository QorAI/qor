"""
QOR Documents — PDF/DOCX Text & Image Extraction
===================================================
Extracts text and images from document files for the QOR knowledge pipeline.

Supports:
  - PDF: via pymupdf (primary) or pdfplumber (fallback)
  - DOCX: via python-docx

Usage as a tool:
    result = read_document("report.pdf")

All dependencies are optional — graceful fallback with install hints.
"""

import os
import logging
from typing import List, Optional

logger = logging.getLogger(__name__)

# Standard search directories for document files
_SEARCH_DIRS = [".", "data", "knowledge", "documents", "qor-data",
                os.path.join("qor-data", "knowledge")]

# Maximum text characters to return (prevents huge payloads)
_MAX_TEXT_CHARS = 4000


def _find_file(filename: str) -> Optional[str]:
    """Search standard directories for a file."""
    # If absolute or relative path exists, use it directly
    if os.path.isfile(filename):
        return filename
    # Search common directories
    basename = os.path.basename(filename)
    for d in _SEARCH_DIRS:
        candidate = os.path.join(d, basename)
        if os.path.isfile(candidate):
            return candidate
    return None


def _extract_pdf_text(path: str) -> str:
    """Extract text from PDF using pymupdf (primary) or pdfplumber (fallback)."""
    # Try pymupdf first (fast, widely used)
    try:
        import pymupdf
        doc = pymupdf.open(path)
        pages = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                pages.append(text.strip())
        doc.close()
        return "\n\n".join(pages)
    except ImportError:
        pass

    # Fallback: pdfplumber
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    pages.append(text.strip())
        return "\n\n".join(pages)
    except ImportError:
        pass

    return ""


def _extract_docx_text(path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        import docx
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except ImportError:
        return ""


def extract_document_images(path: str) -> list:
    """
    Extract embedded images from a document.

    Args:
        path: Path to PDF or DOCX file

    Returns:
        List of PIL.Image objects (empty list if extraction fails or no images)
    """
    images = []
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        try:
            import pymupdf
            from PIL import Image
            import io

            doc = pymupdf.open(path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    if base_image:
                        image_bytes = base_image["image"]
                        img = Image.open(io.BytesIO(image_bytes))
                        images.append(img)
            doc.close()
        except (ImportError, Exception) as e:
            logger.debug(f"Image extraction from PDF failed: {e}")

    elif ext == ".docx":
        try:
            import docx
            from PIL import Image
            import io

            doc = docx.Document(path)
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_bytes = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_bytes))
                    images.append(img)
        except (ImportError, Exception) as e:
            logger.debug(f"Image extraction from DOCX failed: {e}")

    return images


def read_document(query: str) -> str:
    """
    Read text from a PDF or DOCX document.

    Tool handler for the QOR tool system.

    Args:
        query: Filename or path to document (e.g. "report.pdf", "notes.docx")

    Returns:
        Extracted text (truncated to 4000 chars) or error message
    """
    # Extract filename from query (strip common prefixes)
    filename = query.strip()
    for prefix in ["read ", "open ", "load ", "extract "]:
        if filename.lower().startswith(prefix):
            filename = filename[len(prefix):].strip()

    # Detect file type
    ext = os.path.splitext(filename)[1].lower()
    if ext not in (".pdf", ".docx"):
        # Try to guess from query content
        if "pdf" in query.lower():
            ext = ".pdf"
        elif "docx" in query.lower() or "doc" in query.lower():
            ext = ".docx"
        else:
            return f"Unsupported document type: {ext or 'unknown'}. Supported: .pdf, .docx"

    # Find the file
    path = _find_file(filename)
    if not path:
        return f"Document not found: {filename}. Searched: {', '.join(_SEARCH_DIRS)}"

    # Extract text
    text = ""
    if ext == ".pdf":
        text = _extract_pdf_text(path)
        if not text:
            return (
                f"Could not extract text from {filename}. "
                "Install a PDF library: pip install pymupdf  (or: pip install pdfplumber)"
            )
    elif ext == ".docx":
        text = _extract_docx_text(path)
        if not text:
            return (
                f"Could not extract text from {filename}. "
                "Install: pip install python-docx"
            )

    # Truncate if needed
    if len(text) > _MAX_TEXT_CHARS:
        text = text[:_MAX_TEXT_CHARS] + f"\n\n[Truncated — showing first {_MAX_TEXT_CHARS} of {len(text)} characters]"

    return f"Document: {os.path.basename(path)}\n\n{text}"
